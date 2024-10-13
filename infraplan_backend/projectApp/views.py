from rest_framework import generics, status
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from .models import Project
from .serializers import ProjectSerializer
from django.core.exceptions import ValidationError
from docx import Document


# Create a new project
class ProjectCreateView(generics.CreateAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


# List all projects
class ProjectListView(generics.ListAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]


# Retrieve, update, or delete a specific project (combined view)
class ProjectDetailView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        try:
            self.perform_update(serializer)
            return Response(serializer.data)
        except ValidationError as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)


# Upload a .docx file and extract content
class ProjectUploadView(generics.GenericAPIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        file = request.FILES.get('file')
        if file and file.name.endswith('.docx'):
            document = Document(file)
            details = ""
            for para in document.paragraphs:
                details += para.text + "\n"
            return Response({"details": details}, status=status.HTTP_200_OK)
        return Response({"error": "Invalid file format. Please upload a .docx file."}, status=status.HTTP_400_BAD_REQUEST)


# Get all projects filtered by status
class ProjectByStatusView(generics.ListAPIView):
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        status_filter = self.request.query_params.get('status', None)
        if status_filter:
            return Project.objects.filter(status=status_filter)
        return Project.objects.none()  # Return an empty queryset if no status is provided


# Get all projects filtered by field
class ProjectByFieldView(generics.ListAPIView):
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        field_filter = self.request.query_params.get('field', None)
        if field_filter:
            return Project.objects.filter(field__icontains=field_filter)
        return Project.objects.none()  # Return an empty queryset if no field is provided


# Independent View to Delete a specific project
class ProjectDeleteView(generics.DestroyAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def delete(self, request, *args, **kwargs):
        instance = self.get_object()
        instance.delete()
        return Response({"message": "Project deleted successfully."}, status=status.HTTP_204_NO_CONTENT)


# Independent View to Update a specific project
class ProjectUpdateView(generics.UpdateAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def update(self, request, *args, **kwargs):
        instance = self.get_object()

        # Extract file if uploaded
        file = request.FILES.get('certificate')  # Adjust based on the field name for the file upload
        if file and file.name.endswith('.docx'):
            document = Document(file)
            file_content = "\n".join([para.text for para in document.paragraphs])
            # Use file content as description if file is uploaded
            request.data['description'] = file_content

        # Validate and update the project
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        try:
            self.perform_update(serializer)
            return Response(serializer.data)
        except ValidationError as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)